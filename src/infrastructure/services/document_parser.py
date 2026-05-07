"""Service for parsing different document formats and extracting text."""

from typing import Dict, Any, List, Optional
from pydantic import BaseModel, Field
import io

class ParsedDocument(BaseModel):
    """Result of document parsing."""
    file_name: str
    file_type: str
    text: str = ""
    error: Optional[str] = None
    has_tables: bool = False
    table_count: int = 0
    word_count: int = 0
    tables: List[List[Any]] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)
    entities: Dict[str, Any] = Field(default_factory=dict)

    def get_tables_as_dict(self) -> List[List[Any]]:
        """Return tables extracted during parsing."""
        return self.tables


class DocumentParserService:
    """Service for parsing documents to extract text and tables."""
    
    async def parse_document(
        self, file_content: bytes, file_name: str, file_type: str
    ) -> ParsedDocument:
        """Parse document content based on file type.
        
        Args:
            file_content: Binary file content
            file_name: Name of the file
            file_type: Extension or type of the file
            
        Returns:
            ParsedDocument containing text and metadata
        """
        doc = ParsedDocument(file_name=file_name, file_type=file_type)
        try:
            ftype = file_type.lower()
            if ftype.endswith('.pdf') or ftype == 'pdf':
                import pdfplumber
                text_parts = []
                all_tables: list = []
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text_parts.append(extracted)
                        t = page.extract_table()
                        if t:
                            all_tables.append(t)
                doc.text = "\n".join(text_parts)
                doc.tables = all_tables
                doc.table_count = len(all_tables)
                doc.has_tables = bool(all_tables)
                doc.word_count = len(doc.text.split())
            elif ftype.endswith('.docx') or ftype == 'docx':
                from docx import Document
                document = Document(io.BytesIO(file_content))
                doc.text = "\n".join([p.text for p in document.paragraphs])
                doc.word_count = len(doc.text.split())
            elif ftype.endswith('.xlsx') or ftype == 'xlsx':
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
                    text_parts = []
                    all_tables = []
                    for sheet in wb.worksheets:
                        rows = list(sheet.iter_rows(values_only=True))
                        if not rows:
                            continue
                        # First row as headers, rest as data
                        headers = [str(c) if c is not None else '' for c in rows[0]]
                        data_rows = []
                        for row in rows[1:]:
                            row_vals = [str(c) if c is not None else '' for c in row]
                            if any(v.strip() for v in row_vals):
                                data_rows.append(dict(zip(headers, row_vals)))
                                text_parts.append('\t'.join(row_vals))
                        if data_rows:
                            all_tables.append({'headers': headers, 'rows': data_rows[:50]})
                    doc.text = '\n'.join(text_parts)
                    doc.tables = all_tables
                    doc.table_count = len(all_tables)
                    doc.has_tables = bool(all_tables)
                    doc.word_count = len(doc.text.split())
                except ImportError:
                    doc.error = "openpyxl not installed — cannot parse .xlsx files"
            elif ftype.endswith('.txt') or ftype.endswith('.csv') or ftype in ['txt', 'csv']:
                doc.text = file_content.decode('utf-8', errors='ignore')
                doc.word_count = len(doc.text.split())
            else:
                doc.error = f"Unsupported file type for parsing: {file_type}"
                
        except Exception as e:
            doc.error = str(e)
            
        return doc
